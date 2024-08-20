#!/usr/bin/python3
import os, sys#, stat
import time
import subprocess
from subprocess import Popen, PIPE
from xml.sax.saxutils import escape as xml_escape
import re
import json
import openpyxl
import hashlib
import textract

first = lambda x: x[0]
second = lambda x: x[1]

hostname = "http://195.148.30.97"
kp_hostname = "https://kielipankki.fi"
wrkdir = "/var/www/kielipankki-tools"
path_to_tagtools = "/usr/local/bin/"

kpdemos_header = '''
<header>
<a href="http://kielipankki.fi/tools/demo/">Back to main demo page</a>
</header>
'''

def mx_auto(t):
    return wrap_in_tags(wrap_in_tags(t, 'div', attribs='class="col center-block"'), 'div', attribs='class="row mx-auto"')

def add_attribute(key, value, xml_tag):
    if xml_tag.startswith('<') and xml_tag.endswith('>'):
        return xml_tag[:-1] + " " + key + '="' + xml_escape(value) + '">'
    else:
        return xml_tag

def log(message):
     logfilename = os.path.basename(sys.argv[0] + '.log')
     if logfilename == "":
         logfilename = "misc.log"
     localtime = time.localtime()
     path = os.path.join(wrkdir, "log", logfilename)
     timestring = "{year}-{month}-{day} {hour}:{minute}:{second}".format(
         year = localtime.tm_year,
         month = localtime.tm_mon,
         day = localtime.tm_mday,
         hour = localtime.tm_hour,
         minute = localtime.tm_min,
         second = localtime.tm_sec)
     writefile = open(path, "a", encoding = "utf-8")
     writefile.write(timestring + " " + message.strip() + "\n")
     writefile.close()
# #    st = os.stat(path)
# #    os.chmod(path, st.st_mode | stat.S_IWOTH)

def tokenize(text):
    process = Popen([path_to_tagtools + "finnish-tokenize"], stdin=PIPE, stdout=PIPE, stderr=PIPE)
    out, err = process.communicate(input=text.replace('\n', '\n\n').encode("utf-8"))
    out_utf8 = out[:-1].decode("utf-8")
    retval = []
    thissent = []
    for line in out_utf8.split("\n"):
        if line == "":
            if len(thissent) > 0:
                retval.append(thissent)
            thissent = []
            continue
        thissent.append(line)
#    retval.append([text.replace(' .\n', ' .\n\n').replace('\n', '#')])
    return retval

def clean_tempfiles():
    subprocess.call(["/var/www/kielipankki-tools/clean-tmpfiles"])

def text_from_file(form_file):
    hashcode = hashlib.sha1(str(time.time()).encode("utf-8")).hexdigest()
    if '.' not in form_file.filename:
        filename = form_file.filename
        ext = '.txt'
    else:
        filename = form_file.filename[:form_file.filename.rindex('.')]
        ext = form_file.filename[form_file.filename.rindex('.'):]
    if ext == '.txt':
        file_contents = form_file.file.read()
        try:
            file_contents.decode("utf-8")
        except UnicodeDecodeError:
            file_content = file_contents.decode("latin1").encode("utf-8")
        return file_contents
    uploaded_file_path = os.path.join(wrkdir + "/tmp/", "upload_" + hashcode + ext)
    with open(uploaded_file_path, "wb") as f:
        f.write(form_file.file.read())
    if ext in ('.png', '.jpg', '.jpeg', '.gif', '.tiff'):
        retval = textract.process(uploaded_file_path, encoding="utf-8", extension = ext[1:], method="tesseract", language="fin")
    else:
        retval = textract.process(uploaded_file_path, encoding="utf-8", extension = ext[1:])
    retval = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', str(retval, encoding = "utf-8")).encode("utf-8")
    os.remove(uploaded_file_path)
    return retval
    
def text_from_text_file(form_file):
    file_contents = form_file.file.read()
    try:
        file_contents = file_contents.decode("utf-8")
    except UnicodeDecodeError:
        file_content = file_contents.decode("latin1")
    return file_contents
    
def ocr_from_file(form_file, lang):
    hashcode = hashlib.sha1(str(time.time()).encode("utf-8")).hexdigest()
    if '.' not in form_file.filename:
        filename = form_file.filename
        ext = '.txt'
    else:
        filename = form_file.filename[:form_file.filename.rindex('.')]
        ext = form_file.filename[form_file.filename.rindex('.'):]
    if ext == '.txt':
        file_contents = form_file.file.read()
        try:
            file_contents.decode("utf-8")
        except UnicodeDecodeError:
            file_content = file_contents.decode("latin1").encode("utf-8")
        return file_contents, hashcode, ''
    uploaded_file_path = os.path.join(wrkdir + '/tmp/upload_' + hashcode + ext)
    with open(uploaded_file_path, "wb") as f:
        f.write(form_file.file.read())
    if ext in ('.png', '.jpg', '.jpeg', '.gif', 'pdf', 'svg', 'tiff'):
        retval = textract.process(uploaded_file_path, encoding="utf-8", extension = ext[1:], method="tesseract", language=lang)
    else:
        retval = textract.process(uploaded_file_path, encoding="utf-8", extension = ext[1:])
    retval = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', str(retval, encoding = "utf-8")).encode("utf-8")
        #    os.remove(uploaded_file_path)
    return retval, hashcode, "upload_" + hashcode + ext
    
def make_doctype():
    return "<!doctype html>\n" #"Content-type: text/html\n\n<!doctype html>\n"

def text_to_html(text):
    retval = ''
    paras = text.split('\n\n')
    for para in paras:
        retval += '<p>' + para.replace('\n', '<br/>') + '</p>'
    return wrap_in_tags(retval, 'p')

def pad_rows(rows, upto = None, pad_with = ''):
    if upto == None:
        maxlen = max(map(len, rows))
    else:
        maxlen = upto
    return [row if len(row) >= maxlen else row + ((maxlen - len(row)) * [pad_with]) for row in rows]

def extract_column(rows, n):
    retval = []
    for row in rows:
        retval.append(row[n])
    return retval

def tsv2rows(tsv):
    return pad_rows(list(map(lambda x: x.split('\t'), tsv.split('\n'))))

def cols2rows(cols, pad_with = ''):
    maxlen = max(map(len, cols))
    retval = []
    for i in range(maxlen):
        this_row = []
        for col in cols:
            if len(col) <= i:
                this_row.append(pad_with)
            else:
                this_row.append(col[i])
        retval.append(this_row)
    return retval

def paste_new_column(rows, new_column, pad = " "):
    retval = []
    for i, row in enumerate(rows):
        if len(new_column) >= i:
            retval.append(row + [new_column[i]])
        else:
            retval.append(row + [pad])
    return retval

def make_table(rows, header = [], tdattribs = ""):

    def nice_format(item):
        if type(item) == type(float()):
            return '{:.3f}'.format(item)
        return str(item)

    def escape(s):
        if s.startswith('<img src="' + hostname) or s.startswith('<img src="' + kp_hostname):
            return s
        return xml_escape(s)
    
    retval = ''.join(list(map(lambda x: wrap_in_tags(nice_format(x), 'th'), header)))
    for row in rows:
        this_row = ''
        for item in row:
            this_row += wrap_in_tags(escape(nice_format(item)), "td", attribs=tdattribs.format(CELL=item))
        retval += wrap_in_tags(this_row, "tr", oneline = False)
    return wrap_in_tags(retval, "table", oneline = False, attribs = "class=table")

def make_tsv(rows):
    return '\n'.join(map('\t'.join, rows))
rows2tsv = make_tsv

def write_tsv(tsv, session_key, dest_dir=None):
    if not dest_dir:
        dest_dir = wrkdir + "/tmp/"
    with open(dest_dir + session_key + ".tsv", "w", encoding = 'utf-8') as f:
        f.write(tsv)

def write_txt(txt, session_key):
    with open(wrkdir + "/tmp/" + session_key + ".txt", "w", encoding = 'utf-8') as f:
        f.write(txt + "\n")

def write_file(data, extension, session_key):
    with open(wrkdir + "/tmp/" + session_key + "." + extension, "w", encoding = 'utf-8') as f:
        f.write(data)

def abbreviate_text(text, n = 60):
    if len(text) <= n:
        return text
    return text[:n-3] + '...'
        
def permute_rows(rows, p):
    max_slot = max(p) + 1
    padded_rows = pad_rows(rows, upto=max_slot)
    retval = []
    for row in padded_rows:
        this_row = []
        for slot in p:
            this_row.append(row[slot])
        retval.append(this_row)
    return retval

def permute_tsv(s, p):
    return make_tsv(permute_rows(tsv2rows(s), p))

def make_head(title = "Untitled demo", scripts = (), extra_meta = ""):
    retval = wrap_in_tags(title, "title")
    retval += '<meta charset="utf-8"/>\n'
    retval += extra_meta
    retval += '<script src="https://code.jquery.com/jquery-3.3.1.min.js"></script>\n'
    retval += '<script src="https://unpkg.com/popper.js/dist/umd/popper.min.js"></script>\n'
    retval += '<script src="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/js/bootstrap.min.js"></script>\n'
    retval += '<script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/2.7.3/Chart.bundle.min.js"></script>\n'
    retval += '<link href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" rel="stylesheet" />\n'
    retval += "<style> .bold { font-weight: bold; } </style>"
    retval += '''<style> canvas{
    -moz-user-select: none;
    -webkit-user-select: none;
    -ms-user-select: none;
    }

    p{
    word-wrap: break-word;
    }
    </style>
'''
    retval += '''
    <link rel="icon" type="image/png" sizes="32x32" href="resources/favicon-32x32.png">
    <link rel="icon" type="image/png" sizes="96x96" href="resources/favicon-96x96.png">
    <link rel="icon" type="image/png" sizes="16x16" href="resources/favicon-16x16.png">
    '''
    for script in scripts:
        retval += wrap_in_tags(script, "script", oneline = False)
    return wrap_in_tags(retval, "head", oneline = False)

def wrap_in_tags(content, tag, oneline = True, attribs = None):
    if not attribs:
        attribs = ""
    else:
        attribs = " " + attribs
    if oneline:
        return "<" + tag + attribs + ">" + content + "</" + tag + ">\n"
    else:
        return "<" + tag + attribs + ">\n" + content + "\n</" + tag + ">\n"

def wrap_html(head, body, onload = None):
    if onload:
        return make_doctype() + "<html>\n" + head + '<body onload="{}">\n'.format(onload) + kpdemos_header + body + "</body>\n</html>\n"
    return make_doctype() + "<html>\n" + head + '<body>\n' + kpdemos_header + body + "</body>\n</html>\n"

def write_excel(rows, filename, title):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title
    if len(rows) > 0 and rows[0][-1] == 'Morphology':
        ws.append(rows[0])
        for row in rows[1:]:
            morpho = row[-1]
            parts = morpho.replace('\t', '').replace('][', '\t').replace('[', '').replace(']', '').split('\t')
            POS = ''
            NUM = ''
            CASE = ''
            VOICE = ''
            MOOD = ''
            TENSE = ''
            PERS = ''
            OTHER = []
            for part in parts:
                attrib = part.split('=')[0]
                if attrib == 'POS':
                    POS = part
                elif attrib == 'NUM':
                    NUM = part
                elif attrib == 'CASE':
                    CASE = part
                elif attrib == 'VOICE':
                    VOICE = part
                elif attrib == 'MOOD':
                    MOOD = part
                elif attrib == 'TENSE':
                    TENSE = part
                elif attrib == 'PERS':
                    PERS = part
                else:
                    OTHER.append(part)
            this_row = row[:-1]
            this_row.append(POS)
            this_row.append(NUM)
            this_row.append(CASE)
            this_row.append(VOICE)
            this_row.append(MOOD)
            this_row.append(TENSE)
            this_row.append(PERS)
            this_row.append('|'.join(OTHER))
            ws.append(this_row)
    else:
        for row in rows:
            ws.append(row)
    for column_cells in ws.columns:
        length = min(max(len(str(cell.value)) for cell in column_cells), 20)
        ws.column_dimensions[column_cells[0].column].width = length + 3
    wb.save(wrkdir + "/tmp/" + filename + ".xlsx")

def write_docx(txt, session_key, title):
    import docx
    document = docx.Document()
    document.add_heading(title, 0)
    document.add_paragraph(re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', txt))
    #document.add_paragraph(re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', txt))
    document.save(wrkdir + "/tmp/" + session_key + ".docx")

kielipankki_texttools_api_url = 'http://kielipankki.rahtiapp.fi/text/fi'
import requests
import json
    
def postag(text):
    submit_url = kielipankki_texttools_api_url + '/postag'
    
def sentiment(text):
    submit_url = kielipankki_texttools_api_url + '/sentiment'

def nertag(text, show_analyses = False):
    submit_url = kielipankki_texttools_api_url + '/nertag/submit'
    query_url = kielipankki_texttools_api_url + '/nertag/query_job'
    if type(text) == bytes:
        encoded_text = text
    else:
        encoded_text = text.encode('utf-8')
    response = requests.post(submit_url,
                             data = encoded_text,
                             params = {'show-analyses': show_analyses})
    response.raise_for_status()
    jobid = json.loads(response.text).get('jobid')
    while True:
        time.sleep(2)
        response = requests.post(query_url,
                                 data = jobid)
        response_dict = json.loads(response.text)
        if response_dict.get('status') == 'pending':
            continue
        else:
            return response_dict['result']
